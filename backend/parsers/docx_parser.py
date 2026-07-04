"""
DOCX (Microsoft Word) text extractor.
Extracts text from .docx binary content using python-docx,
then hands the plain text to the main ScriptParser cascade.
"""
from typing import Optional


def extract_docx_text(file_bytes: bytes) -> str:
    """
    Extract all paragraph text from a .docx file.
    Returns a plain-text string suitable for ScriptParser.parse().
    Raises ImportError if python-docx is not installed.
    Raises ValueError on corrupt/invalid files.
    """
    try:
        import docx  # python-docx
    except ImportError:
        raise ImportError(
            "python-docx is required to read .docx files. "
            "Install it with: pip install python-docx"
        )

    import io
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Cannot open .docx file: {e}")

    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Also extract text from tables (common in A/V scripts)
    for table in doc.tables:
        for row in table.rows:
            row_parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_parts:
                lines.append("  ".join(row_parts))

    return "\n".join(lines)
